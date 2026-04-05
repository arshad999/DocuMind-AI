from pathlib import Path
from io import BytesIO
from typing import List
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


def ensure_directory(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def safe_filename(filename: str) -> str:
    return Path(filename).name.replace("..", "").replace("/", "_").replace("\\", "_")


def text_to_docx(text: str, output_path: Path) -> None:
    doc = Document()
    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        else:
            doc.add_paragraph(stripped)
    doc.save(output_path)


def text_to_pdf(text: str, output_path: Path) -> None:
    buffer = BytesIO()
    doc = canvas.Canvas(str(output_path), pagesize=letter)
    width, height = letter
    y = height - 50
    for paragraph in text.split("\n\n"):
        for line in paragraph.splitlines():
            if y < 80:
                doc.showPage()
                y = height - 50
            doc.setFont("Helvetica", 11)
            doc.drawString(50, y, line[:95])
            y -= 15
        y -= 10
    doc.save()
